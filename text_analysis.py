import streamlit as st
import matplotlib.pyplot as plt
import pandas as pd
import seaborn as sns
from transformers import pipeline, AutoTokenizer, AutoModelForSequenceClassification
from wordcloud import WordCloud
import emoji
import torch
import os
from pathlib import Path
import logging
import traceback
import re
from collections import defaultdict
from docx import Document
import PyPDF2
import io
import base64
import tempfile

# Define detailed sentiment categories
SENTIMENT_CATEGORIES = {
    "JOYFUL": {"emoji": "😊", "description": "Expressing great happiness"},
    "EXCITED": {"emoji": "🤩", "description": "Showing enthusiasm and eagerness"},
    "LOVING": {"emoji": "❤️", "description": "Showing love or affection"},
    "GRATEFUL": {"emoji": "🙏", "description": "Showing appreciation"},
    "HOPEFUL": {"emoji": "🌟", "description": "Showing hope"},
    "PROUD": {"emoji": "🦚", "description": "Showing satisfaction from achievements"},
    "PEACEFUL": {"emoji": "🕊️", "description": "Tranquil and undisturbed"},
    "SAD": {"emoji": "😢", "description": "Showing sorrow"},
    "ANGRY": {"emoji": "😠", "description": "Showing annoyance or hostility"},
    "AFRAID": {"emoji": "😨", "description": "Showing fear"},
    "GUILTY": {"emoji": "😣", "description": "Showing guilt"},
    "ASHAMED": {"emoji": "😳", "description": "Showing shame"},
    "LONELY": {"emoji": "🔕", "description": "Showing loneliness"},
    "STRESSED": {"emoji": "😫", "description": "Showing stress"},
    "SURPRISED": {"emoji": "😲", "description": "Showing surprise"},
    "CURIOUS": {"emoji": "🤔", "description": "Eager to learn"},
    "THOUGHTFUL": {"emoji": "💭", "description": "Showing consideration"},
    "FOCUSED": {"emoji": "🎯", "description": "Showing attention"},
    "CALM": {"emoji": "😌", "description": "Showing peace"},
    "TIRED": {"emoji": "😴", "description": "Showing weariness"},
    "BITTERSWEET": {"emoji": "🥲", "description": "Mixed sadness and happiness"},
    "NOSTALGIC": {"emoji": "💫", "description": "Sentimental longing"},
    "MELANCHOLY": {"emoji": "🌧️", "description": "Pensive sadness"},
    "EMPATHETIC": {"emoji": "🤝", "description": "Understanding others' feelings"},
    "RESENTFUL": {"emoji": "😤", "description": "Showing bitterness"}
}

def setup_logging():
    logging.basicConfig(level=logging.INFO)
    return logging.getLogger(__name__)

def initialize_model(model_name, logger, cache_path):
    try:
        cache_dir = Path(cache_path)
        cache_dir.mkdir(parents=True, exist_ok=True)
        
        logger.info(f"Using cache directory: {cache_dir}")
        logger.info(f"Attempting to load model: {model_name}")
        
        tokenizer = AutoTokenizer.from_pretrained(model_name, cache_dir=cache_dir)
        model = AutoModelForSequenceClassification.from_pretrained(model_name, cache_dir=cache_dir)
        
        sentiment_pipeline = pipeline(
            "sentiment-analysis",
            model=model,
            tokenizer=tokenizer,
            device=0 if torch.cuda.is_available() else -1
        )
        
        logger.info("Model loaded successfully")
        return sentiment_pipeline
    
    except Exception as e:
        logger.error(f"Error initializing model: {str(e)}\n{traceback.format_exc()}")
        st.error(f"An error occurred while loading the model: {str(e)}")
        return None

def map_to_detailed_sentiment(basic_sentiment, confidence):
    # Map basic sentiment to detailed categories based on confidence and patterns
    # This is a simple example - you might want to use a more sophisticated mapping
    if basic_sentiment == "POSITIVE":
        if confidence > 0.9:
            return "JOYFUL"
        elif confidence > 0.8:
            return "EXCITED"
        elif confidence > 0.7:
            return "HOPEFUL"
        else:
            return "PEACEFUL"
    elif basic_sentiment == "NEGATIVE":
        if confidence > 0.9:
            return "ANGRY"
        elif confidence > 0.8:
            return "SAD"
        elif confidence > 0.7:
            return "AFRAID"
        else:
            return "STRESSED"
    else:
        return "THOUGHTFUL"

def split_text(text):
    return re.split(r'[.,!?;:]', text)

def generate_pdf(text):
    pdf_bytes = io.BytesIO()
    doc = Document()
    doc.add_paragraph(text)
    doc.save(pdf_bytes)
    pdf_bytes.seek(0)
    return pdf_bytes.getvalue()

def detect_sentiments(sentences, sentiment_pipeline, detailed=False, use_detailed_categories=False):
    categorized_sentiments = defaultdict(list)
    results = sentiment_pipeline(sentences)
    
    sentiments = []
    confidence_scores = []
    report = []
    
    for text, result in zip(sentences, results):
        base_sentiment = result['label']
        confidence = result['score']
        
        if use_detailed_categories:
            sentiment = map_to_detailed_sentiment(base_sentiment, confidence)
        else:
            sentiment = base_sentiment
            
        categorized_sentiments[sentiment].append(text.strip())
        sentiments.append(sentiment)
        confidence_scores.append(confidence)
        
        segment_report = [
            f"**Segment:** {text.strip()}",
            f"Sentiment: {sentiment} {SENTIMENT_CATEGORIES.get(sentiment, {}).get('emoji', '🤔')}",
            f"Confidence: {confidence:.2f}"
        ]
        
        if detailed:
            word_results = []
            for word in text.split():
                word_result = sentiment_pipeline([word])[0]
                word_sentiment = map_to_detailed_sentiment(word_result['label'], word_result['score']) if use_detailed_categories else word_result['label']
                word_results.append(f"- **{word}**: {word_sentiment} {SENTIMENT_CATEGORIES.get(word_sentiment, {}).get('emoji', '🤔')} ({word_result['score']:.2f})")
            segment_report.extend(word_results)
        
        report.extend(segment_report + ["---"])
    
    # Visualizations
    st.subheader("Overall Sentiment Analysis")
    
    df = pd.DataFrame({"Sentiment": sentiments, "Confidence": confidence_scores})
    
    # Sentiment distribution plot
    fig, ax = plt.subplots(figsize=(10, 6))
    sns.countplot(data=df, x="Sentiment", palette="coolwarm", ax=ax)
    plt.xticks(rotation=45)
    st.pyplot(fig)
    
    # Confidence distribution plot
    fig, ax = plt.subplots(figsize=(10, 6))
    sns.histplot(df["Confidence"], bins=10, kde=True, ax=ax)
    plt.title("Confidence Score Distribution")
    st.pyplot(fig)
    
    # Word cloud
    if len(" ".join(sentences)) > 0:
        st.subheader("Word Cloud")
        wordcloud = WordCloud(width=800, height=400, background_color='white').generate(" ".join(sentences))
        fig, ax = plt.subplots(figsize=(10, 6))
        ax.imshow(wordcloud, interpolation='bilinear')
        ax.axis("off")
        st.pyplot(fig)
    
    return "\n".join(report)

def analyze_file(uploaded_file, sentiment_pipeline, detailed=False, use_detailed_categories=False):
    if uploaded_file.type == "application/pdf":
        pdf_reader = PyPDF2.PdfReader(uploaded_file)
        text = "\n".join([page.extract_text() for page in pdf_reader.pages if page.extract_text()])
    elif uploaded_file.type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
        doc = Document(uploaded_file)
        text = "\n".join([para.text for para in doc.paragraphs])
    elif uploaded_file.type == "text/plain":
        text = uploaded_file.getvalue().decode("utf-8")
    else:
        st.error("Unsupported file type")
        return
    
    sentences = [s.strip() for s in split_text(text) if s.strip()]
    return detect_sentiments(sentences, sentiment_pipeline, detailed, use_detailed_categories)

def run():
    logger = setup_logging()
    st.header("📖 Text Sentiment Analysis")
    
    default_cache = os.path.join(os.path.expanduser("~"), ".cache", "huggingface")
    cache_path = st.text_input("Cache Directory Path:", value=default_cache)
    
    # Model selection
    model_options = {
        "DistilBERT (Fast)": "distilbert-base-uncased-finetuned-sst-2-english",
        "BERT": "nlptown/bert-base-multilingual-uncased-sentiment",
        "RoBERTa": "cardiffnlp/twitter-roberta-base-sentiment",
        "GPT2 for Emotions": "microsoft/DialogRPT-updown",
        "bhadresh-savani":"bhadresh-savani/distilbert-base-uncased-emotion"  # Note: This is an example, verify compatibility
    }
    
    selected_model = st.selectbox("Choose a sentiment analysis model:", list(model_options.keys()))
    sentiment_pipeline = initialize_model(model_options[selected_model], logger, cache_path)
    
    if sentiment_pipeline is None:
        st.warning("Please try selecting a different model or check your internet connection.")
        return
    
    # Analysis options
    st.subheader("Analysis Options")
    use_detailed_categories = st.checkbox("Use detailed sentiment categories")
    show_word_analysis = st.checkbox("Enable word-level analysis")
    
    # Input methods
    user_input = st.text_area("Enter text for sentiment analysis:")
    uploaded_file = st.file_uploader("Upload a text, PDF, or DOCX file", type=["txt", "pdf", "docx"])
    
    if st.button("Analyze"):
        report = None
        
        if uploaded_file:
            report = analyze_file(uploaded_file, sentiment_pipeline, show_word_analysis, use_detailed_categories)
        elif user_input:
            sentences = [s.strip() for s in split_text(user_input) if s.strip()]
            report = detect_sentiments(sentences, sentiment_pipeline, show_word_analysis, use_detailed_categories)
        
        if report:
            st.markdown(report)
            pdf_data = generate_pdf(report)
            st.download_button(
                label="Download Analysis as PDF",
                data=pdf_data,
                file_name="sentiment_analysis.pdf",
                mime="application/pdf"
            )

if __name__ == "__main__":
    run()