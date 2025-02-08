import streamlit as st
import matplotlib.pyplot as plt
import pandas as pd
import seaborn as sns
from transformers import pipeline, AutoTokenizer, AutoModelForSequenceClassification
from wordcloud import WordCloud
import emoji
import torch
import os
from pathlib import Path
import logging
import traceback
import re
from collections import defaultdict
from docx import Document
import PyPDF2
import io
import base64
import tempfile
from PIL import Image

# Enhanced NLTK imports
import nltk
from nltk.corpus import wordnet as wn
from nltk.corpus import sentiwordnet as swn
from nltk.sentiment import SentimentIntensityAnalyzer
from nltk.tokenize import word_tokenize, sent_tokenize
from nltk.tag import pos_tag
from nltk.stem import WordNetLemmatizer

# Download necessary NLTK data
nltk_packages = [
    'punkt', 'averaged_perceptron_tagger', 'wordnet', 
    'sentiwordnet', 'vader_lexicon', 'omw-1.4'
]

for package in nltk_packages:
    try:
        nltk.data.find(f'tokenizers/{package}')
    except LookupError:
        nltk.download(package)

def setup_logging():
    logging.basicConfig(level=logging.INFO)
    return logging.getLogger(__name__)

def initialize_models(model_name, logger, cache_path):
    try:
        cache_dir = Path(cache_path)
        cache_dir.mkdir(parents=True, exist_ok=True)
        
        logger.info(f"Using cache directory: {cache_dir}")
        logger.info(f"Attempting to load model: {model_name}")
        
        # Initialize transformers model
        tokenizer = AutoTokenizer.from_pretrained(model_name, cache_dir=cache_dir)
        model = AutoModelForSequenceClassification.from_pretrained(model_name, cache_dir=cache_dir)
        
        transformer_pipeline = pipeline(
            "sentiment-analysis",
            model=model,
            tokenizer=tokenizer,
            device=0 if torch.cuda.is_available() else -1
        )
        
        # Initialize NLTK analyzer
        nltk_analyzer = SentimentIntensityAnalyzer()
        
        logger.info("Models loaded successfully")
        return transformer_pipeline, nltk_analyzer
    
    except Exception as e:
        logger.error(f"Error initializing models: {str(e)}\n{traceback.format_exc()}")
        st.error(f"An error occurred while loading the models: {str(e)}")
        return None, None

def get_wordnet_emotion(word):
    """Extract emotion information from WordNet synsets"""
    emotions = set()
    synsets = wn.synsets(word)
    
    # Emotion mapping based on WordNet hierarchies
    emotion_hierarchy = {
        'joy': ['joy', 'happiness', 'pleasure', 'delight'],
        'sadness': ['sadness', 'sorrow', 'grief', 'misery'],
        'anger': ['anger', 'rage', 'fury', 'wrath'],
        'fear': ['fear', 'anxiety', 'worry', 'concern'],
        'surprise': ['surprise', 'amazement', 'astonishment'],
        'disgust': ['disgust', 'revulsion', 'aversion'],
        'trust': ['trust', 'confidence', 'faith'],
        'anticipation': ['anticipation', 'expectation', 'hope']
    }
    
    for synset in synsets:
        # Check definition and lemma names
        definition = synset.definition().lower()
        lemmas = [lemma.name().lower() for lemma in synset.lemmas()]
        
        # Check hypernyms for emotion concepts
        hypernyms = []
        for hypernym in synset.hypernyms():
            hypernyms.extend([lemma.name().lower() for lemma in hypernym.lemmas()])
        
        # Match against emotion categories
        for emotion, keywords in emotion_hierarchy.items():
            for keyword in keywords:
                if (keyword in definition or 
                    keyword in lemmas or 
                    keyword in hypernyms):
                    emotions.add(emotion)
    
    return emotions

def analyze_emotions(text, transformer_pipeline, nltk_analyzer):
    """Comprehensive emotion analysis using multiple approaches"""
    # Tokenize text
    sentences = sent_tokenize(text)
    words = word_tokenize(text)
    pos_tags = pos_tag(words)
    
    # Initialize results
    analysis_results = {
        'transformer': transformer_pipeline([text])[0],
        'nltk_vader': nltk_analyzer.polarity_scores(text),
        'wordnet_emotions': defaultdict(int),
        'detailed_analysis': []
    }
    
    # WordNet emotion analysis
    lemmatizer = WordNetLemmatizer()
    for word, pos in pos_tags:
        # Convert POS tag to WordNet format
        wordnet_pos = {
            'JJ': wn.ADJ, 'VB': wn.VERB,
            'NN': wn.NOUN, 'RB': wn.ADV
        }.get(pos[:2], wn.NOUN)
        
        lemma = lemmatizer.lemmatize(word, wordnet_pos)
        emotions = get_wordnet_emotion(lemma)
        
        for emotion in emotions:
            analysis_results['wordnet_emotions'][emotion] += 1
        
        # Detailed word analysis
        if emotions:
            analysis_results['detailed_analysis'].append({
                'word': word,
                'emotions': emotions,
                'pos': pos,
                'vader_score': nltk_analyzer.polarity_scores(word)
            })
    
    return analysis_results

def generate_visualization(analysis_results):
    """Generate visualizations for the analysis results"""
    figs = []
    
    # Emotion distribution plot
    if analysis_results['wordnet_emotions']:
        fig, ax = plt.subplots(figsize=(10, 6))
        emotions_df = pd.DataFrame(list(analysis_results['wordnet_emotions'].items()),
                                 columns=['Emotion', 'Count'])
        sns.barplot(data=emotions_df, x='Emotion', y='Count', palette='viridis')
        plt.xticks(rotation=45)
        plt.title('Distribution of Emotions (WordNet)')
        figs.append(("Emotion Distribution", fig))
    
    # VADER sentiment scores
    fig, ax = plt.subplots(figsize=(10, 6))
    vader_df = pd.DataFrame([analysis_results['nltk_vader']])
    vader_df.plot(kind='bar', ax=ax)
    plt.title('VADER Sentiment Scores')
    figs.append(("VADER Sentiment", fig))
    
    return figs

def generate_report(text, analysis_results, visualizations):
    """Generate comprehensive analysis report"""
    doc = Document()
    
    # Add title
    doc.add_heading('Emotion Analysis Report', 0)
    
    # Add original text
    doc.add_heading('Analyzed Text', level=1)
    doc.add_paragraph(text)
    
    # Add transformer results
    doc.add_heading('Transformer Model Analysis', level=1)
    doc.add_paragraph(f"Label: {analysis_results['transformer']['label']}")
    doc.add_paragraph(f"Score: {analysis_results['transformer']['score']:.4f}")
    
    # Add VADER results
    doc.add_heading('VADER Sentiment Analysis', level=1)
    for metric, score in analysis_results['nltk_vader'].items():
        doc.add_paragraph(f"{metric}: {score:.4f}")
    
    # Add WordNet emotion analysis
    doc.add_heading('WordNet Emotion Analysis', level=1)
    for emotion, count in analysis_results['wordnet_emotions'].items():
        doc.add_paragraph(f"{emotion}: {count} occurrences")
    
    # Add detailed word analysis
    doc.add_heading('Detailed Word Analysis', level=1)
    for word_analysis in analysis_results['detailed_analysis']:
        p = doc.add_paragraph()
        p.add_run(f"Word: {word_analysis['word']}").bold = True
        p.add_run(f"\nPart of Speech: {word_analysis['pos']}")
        p.add_run(f"\nEmotions: {', '.join(word_analysis['emotions'])}")
        p.add_run(f"\nVADER scores: {word_analysis['vader_score']}")
    
    # Add visualizations
    doc.add_heading('Visualizations', level=1)
    for title, fig in visualizations:
        temp_img = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
        fig.savefig(temp_img.name)
        doc.add_heading(title, level=2)
        doc.add_picture(temp_img.name)
        temp_img.close()
        os.unlink(temp_img.name)
    
    # Save to bytes
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    return doc_bytes.getvalue()

def run():
    logger = setup_logging()
    st.header("📖 Comprehensive Emotion Analysis")
    
    default_cache = os.path.join(os.path.expanduser("~"), ".cache", "huggingface")
    cache_path = st.text_input("Cache Directory Path:", value=default_cache)
    
    # Model selection
    model_options = {
        "RoBERTa (Emotion)": "j-hartmann/emotion-english-distilroberta-base",
        "BERT (Sentiment)": "nlptown/bert-base-multilingual-uncased-sentiment",
        "DistilBERT (Emotion)": "bhadresh-savani/distilbert-base-uncased-emotion"
    }
    
    selected_model = st.selectbox("Choose a model:", list(model_options.keys()))
    transformer_pipeline, nltk_analyzer = initialize_models(model_options[selected_model], logger, cache_path)
    
    if transformer_pipeline is None or nltk_analyzer is None:
        st.warning("Please try selecting a different model or check your internet connection.")
        return
    
    # Input methods
    user_input = st.text_area("Enter text for analysis:")
    uploaded_file = st.file_uploader("Upload a file", type=["txt", "pdf", "docx"])
    
    if st.button("Analyze"):
        text = ""
        if uploaded_file:
            if uploaded_file.type == "application/pdf":
                pdf_reader = PyPDF2.PdfReader(uploaded_file)
                text = "\n".join([page.extract_text() for page in pdf_reader.pages if page.extract_text()])
            elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                doc = Document(uploaded_file)
                text = "\n".join([para.text for para in doc.paragraphs])
            elif uploaded_file.type == "text/plain":
                text = uploaded_file.getvalue().decode("utf-8")
        else:
            text = user_input
        
        if text:
            # Perform analysis
            analysis_results = analyze_emotions(text, transformer_pipeline, nltk_analyzer)
            
            # Generate visualizations
            visualizations = generate_visualization(analysis_results)
            
            # Display visualizations
            for title, fig in visualizations:
                st.subheader(title)
                st.pyplot(fig)
            
            # Generate and offer report download
            report = generate_report(text, analysis_results, visualizations)
            st.download_button(
                label="Download Full Analysis Report",
                data=report,
                file_name="emotion_analysis_report.pdf",
                mime="application/pdf"
            )

if __name__ == "__main__":
    run()