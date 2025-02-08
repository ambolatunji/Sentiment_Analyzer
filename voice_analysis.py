# First, ensure you have all required dependencies by running:
# pip install streamlit matplotlib pandas seaborn transformers wordcloud emoji torch nltk flair numpy python-docx PyPDF2 scikit-learn

import streamlit as st
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import seaborn as sns
from transformers import pipeline, AutoTokenizer, AutoModelForSequenceClassification
from wordcloud import WordCloud
import emoji
import torch
import os
from pathlib import Path
import logging
import traceback
import re
from collections import defaultdict
from docx import Document
import PyPDF2
import io
import base64
import tempfile
import nltk
from nltk.sentiment import SentimentIntensityAnalyzer
from flair.models import TextClassifier
from flair.data import Sentence

def setup_dependencies():
    """Setup and verify all required NLTK downloads"""
    try:
        nltk.data.find('sentiment/vader_lexicon.zip')
    except LookupError:
        nltk.download('vader_lexicon')
    
    try:
        nltk.data.find('tokenizers/punkt')
    except LookupError:
        nltk.download('punkt')

def setup_logging():
    logging.basicConfig(level=logging.INFO)
    return logging.getLogger(__name__)

def initialize_model(model_name, logger, cache_path):
    try:
        cache_dir = Path(cache_path)
        cache_dir.mkdir(parents=True, exist_ok=True)
        
        logger.info(f"Using cache directory: {cache_dir}")
        logger.info(f"Attempting to load model: {model_name}")
        
        tokenizer = AutoTokenizer.from_pretrained(model_name, cache_dir=cache_dir)
        model = AutoModelForSequenceClassification.from_pretrained(model_name, cache_dir=cache_dir)
        
        sentiment_pipeline = pipeline(
            "sentiment-analysis",
            model=model,
            tokenizer=tokenizer,
            device=0 if torch.cuda.is_available() else -1
        )
        
        logger.info("Model loaded successfully")
        return sentiment_pipeline
    
    except Exception as e:
        logger.error(f"Error initializing model: {str(e)}\n{traceback.format_exc()}")
        st.error(f"An error occurred while loading the model: {str(e)}")
        return None

def initialize_analyzers():
    """Initialize NLTK and Flair sentiment analyzers"""
    try:
        sia = SentimentIntensityAnalyzer()
        flair_classifier = TextClassifier.load('en-sentiment')
        return sia, flair_classifier
    except Exception as e:
        st.error(f"Error initializing analyzers: {str(e)}")
        return None, None

def analyze_with_nltk(text, sia):
    try:
        return sia.polarity_scores(text)
    except Exception as e:
        st.error(f"Error in NLTK analysis: {str(e)}")
        return {"compound": 0, "neg": 0, "neu": 0, "pos": 0}

def analyze_with_flair(text, flair_classifier):
    try:
        sentence = Sentence(text)
        flair_classifier.predict(sentence)
        return sentence.labels[0].to_dict()
    except Exception as e:
        st.error(f"Error in Flair analysis: {str(e)}")
        return {"value": "NEUTRAL", "confidence": 0.0}

def split_text(text):
    return [s.strip() for s in re.split(r'[.,!?;:]', text) if s.strip()]

def generate_pdf(text):
    """Generate PDF report"""
    try:
        doc = Document()
        doc.add_heading('Sentiment Analysis Report', 0)
        doc.add_paragraph(text)
        
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes.getvalue()
    except Exception as e:
        st.error(f"Error generating PDF: {str(e)}")
        return None

def detect_sentiments(sentences, sentiment_pipeline, sia, flair_classifier, method="transformer", detailed=False):
    sentiments = []
    confidence_scores = []
    report = []
    
    try:
        for text in sentences:
            if method == "transformer" and sentiment_pipeline:
                result = sentiment_pipeline(text)[0]
                sentiment = result['label']
                confidence = result['score']
            elif method == "nltk" and sia:
                result = analyze_with_nltk(text, sia)
                sentiment = max(result, key=result.get)
                confidence = result[sentiment]
            elif method == "flair" and flair_classifier:
                result = analyze_with_flair(text, flair_classifier)
                sentiment = result['value']
                confidence = result['confidence']
            else:
                continue
            
            sentiments.append(sentiment)
            confidence_scores.append(confidence)
            
            report.extend([
                f"**Segment:** {text}",
                f"Sentiment: {sentiment} (Confidence: {confidence:.2f})",
                ""
            ])
            
            if detailed and method == "transformer" and sentiment_pipeline:
                word_results = sentiment_pipeline(text.split())
                for word, w_result in zip(text.split(), word_results):
                    report.append(f"- **{word}**: {w_result['label']} (Confidence: {w_result['score']:.2f})")
            report.append("---")
        
        # Create visualizations
        if sentiments:
            df = pd.DataFrame({"Sentiment": sentiments, "Confidence": confidence_scores})
            
            # Sentiment distribution
            fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(15, 5))
            
            sns.countplot(data=df, x="Sentiment", palette="coolwarm", ax=ax1)
            ax1.set_title("Sentiment Distribution")
            ax1.tick_params(axis='x', rotation=45)
            
            sns.histplot(data=df, x="Confidence", bins=10, ax=ax2)
            ax2.set_title("Confidence Distribution")
            
            st.pyplot(fig)
            plt.close()
            
            # Word cloud
            if len(" ".join(sentences)) > 0:
                wordcloud = WordCloud(width=800, height=400, background_color='white').generate(" ".join(sentences))
                fig, ax = plt.subplots(figsize=(10, 5))
                ax.imshow(wordcloud, interpolation='bilinear')
                ax.axis("off")
                ax.set_title("Word Cloud")
                st.pyplot(fig)
                plt.close()
        
        return "\n".join(report)
    
    except Exception as e:
        st.error(f"Error in sentiment analysis: {str(e)}")
        return "Error performing sentiment analysis"

def run():
    st.set_page_config(page_title="Sentiment Analysis", layout="wide")
    
    logger = setup_logging()
    setup_dependencies()
    
    st.header("📖 Advanced Sentiment Analysis")
    
    with st.spinner("Initializing analyzers..."):
        sia, flair_classifier = initialize_analyzers()
    
    model_options = {
        "DistilBERT (Fast)": "distilbert-base-uncased-finetuned-sst-2-english",
        "BERT": "nlptown/bert-base-multilingual-uncased-sentiment",
        "RoBERTa": "cardiffnlp/twitter-roberta-base-sentiment",
        "DistilBERT Emotion": "bhadresh-savani/distilbert-base-uncased-emotion"
    }
    
    analysis_methods = {
        "Transformer Model": "transformer",
        "NLTK VADER": "nltk",
        "Flair": "flair"
    }
    
    col1, col2 = st.columns(2)
    
    with col1:
        selected_method = st.selectbox("Choose analysis method:", list(analysis_methods.keys()))
    
    sentiment_pipeline = None
    if selected_method == "Transformer Model":
        with col2:
            selected_model = st.selectbox("Choose a sentiment analysis model:", list(model_options.keys()))
            with st.spinner("Loading transformer model..."):
                sentiment_pipeline = initialize_model(
                    model_options[selected_model],
                    logger,
                    os.path.join(os.path.expanduser("~"), ".cache", "huggingface")
                )
    
    show_word_analysis = st.checkbox("Enable word-level analysis")
    user_input = st.text_area("Enter text for sentiment analysis:", height=150)
    
    if st.button("Analyze"):
        if user_input:
            with st.spinner("Analyzing text..."):
                sentences = split_text(user_input)
                report = detect_sentiments(
                    sentences,
                    sentiment_pipeline,
                    sia,
                    flair_classifier,
                    method=analysis_methods[selected_method],
                    detailed=show_word_analysis
                )
                
                st.markdown(report)
                
                pdf_data = generate_pdf(report)
                if pdf_data:
                    st.download_button(
                        "Download Analysis as PDF",
                        pdf_data,
                        "sentiment_analysis.pdf",
                        "application/pdf"
                    )

if __name__ == "__main__":
    run()